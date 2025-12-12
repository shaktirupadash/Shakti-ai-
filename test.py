import streamlit as st
import google.generativeai as genai
from datetime import datetime
from PIL import Image
import cv2
import tempfile
import time
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
import pandas as pd

# Page config
st.set_page_config(
    page_title="Shakti-Gemini AI Assistant",
    page_icon="⚡",
    layout="wide"
)

# Custom CSS
st.markdown("""
<style>
    .shakti-title {
        font-size: 3.5rem;
        font-weight: bold;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        padding: 20px;
    }
    .subtitle {
        text-align: center;
        color: #888;
        font-size: 1.2rem;
    }
    .sheet-badge {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 5px 15px;
        border-radius: 20px;
        font-weight: bold;
        display: inline-block;
        margin: 10px 0;
    }
</style>
""", unsafe_allow_html=True)

# ---------- Session state init ----------
if "messages" not in st.session_state:
    st.session_state.messages = []

if "current_video_file" not in st.session_state:
    st.session_state.current_video_file = None

if "last_request_time" not in st.session_state:
    st.session_state.last_request_time = 0

if "unit_test_data" not in st.session_state:
    st.session_state.unit_test_data = None

if "uploader_key" not in st.session_state:
    st.session_state.uploader_key = 0

if "clear_uploader" not in st.session_state:
    st.session_state.clear_uploader = False

if "captured_frames" not in st.session_state:
    st.session_state.captured_frames = []

if "excel_sheets" not in st.session_state:
    st.session_state.excel_sheets = {}

if "sheet_analyses" not in st.session_state:
    st.session_state.sheet_analyses = {}

# ----------------------------------------

def smart_rate_limit():
    """Smart rate limiting"""
    current_time = time.time()
    time_since_last = current_time - st.session_state.last_request_time
    
    min_wait = 6
    
    if time_since_last < min_wait:
        wait_time = min_wait - time_since_last
        if wait_time > 0:
            with st.spinner(f"⏳ Rate limiting... waiting {wait_time:.1f}s"):
                time.sleep(wait_time)
    
    st.session_state.last_request_time = time.time()

def generate_unit_test_prompt():
    """Generate a comprehensive prompt for unit test document creation"""
    return """Analyze this video and create a Unit Test Document. For EACH screenshot/frame, provide:

1. A clear, concise description of what action is happening in that frame
2. What the user is doing or what screen is shown
3. Any important text, buttons, or UI elements visible

Format your response as:

# UNIT TEST DOCUMENT

## TEST OVERVIEW
- Test ID: [Generate ID]
- Test Name: [Name based on what you see]
- Feature: [What's being tested]
- Date: [Today's date]

## TEST STEPS

For each frame, write:

**Frame 1:**
[Clear description of what's happening - e.g., "User logged into the system" or "Trip selection screen showing trip 0055600-00"]

**Frame 2:**
[Description of action/screen]

**Frame 3:**
[Description of action/screen]

... continue for all frames

## TEST SUMMARY
- Total Steps: [number of frames]
- Status: Pass/Fail
- Notes: [Any observations]

IMPORTANT: Write clear, simple descriptions for each frame that will go UNDER the screenshot in the document."""

def get_inventory_prompt(analysis_type, sheet_name=""):
    """Get AI prompt based on analysis type"""
    
    sheet_context = f"\n**ANALYZING SHEET: {sheet_name}**\n" if sheet_name else ""
    
    prompts = {
        "stockout_risk": f"""
{sheet_context}
Analyze this data and identify:

1. **CRITICAL STOCKOUTS** (immediate action needed):
   - Products with low stock levels vs. sales velocity
   - Estimated days until stockout
   - Revenue at risk
   - Recommended emergency reorder quantities

2. **WARNING ITEMS** (action needed within 2 weeks):
   - Products approaching low stock
   - Recommended reorder timing and quantities

3. **ADEQUATE STOCK** (monitor):
   - Items with healthy stock levels

For each at-risk item, provide:
- SKU and product name
- Current stock level
- Sales velocity (units per day)
- Estimated stockout date
- Recommended reorder quantity
- Revenue impact

Format as a prioritized action list with clear sections.
""",
        
        "slow_movers": f"""
{sheet_context}
Analyze this data and identify:

1. **EXCESS INVENTORY** (overstock items):
   - Products with very high days of inventory on hand
   - Items with declining sales trends
   - Low turnover rate items
   - Seasonal items past their peak

2. **RECOMMENDATIONS** for each item:
   - Suggested discount percentage for clearance
   - Alternative sales channels
   - Return to vendor options
   - Potential write-off considerations

3. **COST IMPACT**:
   - Calculate total carrying cost of slow-moving inventory
   - Potential savings from clearance actions

Format as an actionable report with specific SKUs and financial impact.
""",
        
        "general_insights": f"""
{sheet_context}
Analyze this data and provide a comprehensive overview:

1. **DATA OVERVIEW**:
   - What type of data is this (inventory, sales, financial, etc.)?
   - Overall assessment of the data health

2. **KEY FINDINGS**:
   - Most important insights from this data
   - Notable patterns or trends
   - Anomalies or outliers

3. **TOP PERFORMERS**:
   - Best performing items/categories
   - Items with optimal metrics

4. **PROBLEM AREAS**:
   - Issues requiring attention
   - Risks or concerns
   - Unusual patterns

5. **RECOMMENDATIONS**:
   - Priority actions
   - Strategic suggestions
   - Specific items requiring attention

Provide clear, business-focused insights with specific examples and actionable recommendations.
"""
    }
    
    return prompts.get(analysis_type, prompts["general_insights"])

def load_excel_sheets(file):
    """Load all sheets from an Excel file"""
    try:
        # Read all sheets
        excel_file = pd.ExcelFile(file)
        sheet_names = excel_file.sheet_names
        
        sheets_dict = {}
        for sheet_name in sheet_names:
            df = pd.read_excel(file, sheet_name=sheet_name)
            sheets_dict[sheet_name] = df
        
        return sheets_dict, None
    except Exception as e:
        return None, str(e)

def analyze_sheet_data(df, sheet_name, analysis_type):
    """Analyze a single sheet's data with Gemini AI"""
    
    # Get basic stats
    total_rows = len(df)
    total_columns = len(df.columns)
    
    # Create data summary
    data_summary = f"""
SHEET: {sheet_name}

DATA SUMMARY:
- Total Rows: {total_rows}
- Total Columns: {total_columns}
- Column Names: {', '.join(df.columns.tolist())}

SAMPLE DATA (First 20 rows):
{df.head(20).to_string()}

BASIC STATISTICS:
{df.describe().to_string()}

DATA TYPES:
{df.dtypes.to_string()}
"""
    
    # Get the analysis prompt
    base_prompt = get_inventory_prompt(analysis_type, sheet_name)
    
    # Create full prompt for AI
    full_prompt = f"""
{base_prompt}

Here's the data to analyze:

{data_summary}

Provide actionable insights and recommendations in a clear, well-structured markdown format.
Use headers, bullet points, and clear sections for easy reading.
"""
    
    return full_prompt

def extract_video_frames_with_save(video_path, num_frames=8):
    """Extract frames from video and save as temporary files"""
    cap = cv2.VideoCapture(video_path)
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    fps = cap.get(cv2.CAP_PROP_FPS)
    
    frames = []
    frame_files = []
    frame_indices = [int(i * total_frames / (num_frames + 1)) for i in range(1, num_frames + 1)]
    
    for i, idx in enumerate(frame_indices):
        cap.set(cv2.CAP_PROP_POS_FRAMES, idx)
        ret, frame = cap.read()
        if ret:
            # Convert BGR to RGB
            frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            pil_image = Image.fromarray(frame_rgb)
            frames.append(pil_image)
            
            # Save frame to temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
            pil_image.save(temp_file.name, 'PNG')
            
            # Calculate timestamp
            timestamp_sec = idx / fps if fps > 0 else 0
            minutes = int(timestamp_sec // 60)
            seconds = int(timestamp_sec % 60)
            
            frame_files.append({
                'file': temp_file.name,
                'index': i + 1,
                'timestamp': f"{minutes:02d}:{seconds:02d}",
                'timestamp_sec': timestamp_sec,
                'frame_number': idx
            })
    
    cap.release()
    return frames, frame_files

def extract_frame_descriptions(test_data, num_frames):
    """Extract descriptions for each frame from the AI response"""
    descriptions = {}
    
    lines = test_data.split('\n')
    current_frame = None
    current_desc = []
    
    for line in lines:
        line = line.strip()
        
        # Look for frame markers
        if line.startswith('**Frame ') and ':' in line:
            # Save previous frame description
            if current_frame and current_desc:
                descriptions[current_frame] = ' '.join(current_desc).strip()
            
            # Start new frame
            try:
                frame_num = int(line.split('Frame ')[1].split(':')[0].strip())
                current_frame = frame_num
                # Get description after the colon
                desc_part = ':'.join(line.split(':')[1:]).strip().strip('*')
                current_desc = [desc_part] if desc_part else []
            except:
                pass
        
        elif current_frame and line and not line.startswith('#') and not line.startswith('**'):
            # Continue description for current frame
            if len(current_desc) < 3:  # Limit description length
                current_desc.append(line)
    
    # Save last frame
    if current_frame and current_desc:
        descriptions[current_frame] = ' '.join(current_desc).strip()
    
    # Fill in missing descriptions
    for i in range(1, num_frames + 1):
        if i not in descriptions:
            descriptions[i] = f"Test execution step {i}"
    
    return descriptions

def create_clean_document_with_images(test_data, frame_files, video_duration=0, video_name="test_video"):
    """Create a clean Word document with large screenshots and descriptions"""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Unit Test Document', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Test ID
    test_id = f"UT-{datetime.now().strftime('%Y%m%d%H%M')}"
    
    # Metadata
    doc.add_paragraph(f'Test ID: {test_id}')
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Video: {video_name}')
    doc.add_paragraph(f'Duration: {int(video_duration//60)}m {int(video_duration%60)}s')
    doc.add_paragraph('_' * 80)
    doc.add_paragraph()
    
    # Extract frame descriptions from AI response
    frame_descriptions = extract_frame_descriptions(test_data, len(frame_files))
    
    # Test Steps with Screenshots
    doc.add_heading('TEST EXECUTION STEPS', level=1)
    doc.add_paragraph('Each screenshot below shows a key step in the test execution.')
    doc.add_paragraph()
    
    for frame_info in frame_files:
        frame_num = frame_info['index']
        
        # Step heading
        step_heading = doc.add_heading(f"Step {frame_num} - Time {frame_info['timestamp']}", level=2)
        
        # Screenshot (LARGE - 6 inches wide for readability)
        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run().add_picture(frame_info['file'], width=Inches(6.0))
        except Exception as e:
            doc.add_paragraph(f"[Screenshot error: {str(e)}]")
        
        doc.add_paragraph()
        
        # Description UNDER the screenshot
        desc_para = doc.add_paragraph()
        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        description = frame_descriptions.get(frame_num, f"Test step {frame_num}")
        
        # Add description in italic, centered
        desc_run = desc_para.add_run(description)
        desc_run.italic = True
        desc_run.font.size = Pt(10)
        desc_run.font.color.rgb = RGBColor(80, 80, 80)
        
        doc.add_paragraph()
        doc.add_paragraph('─' * 80)
        doc.add_paragraph()
    
    # Test Summary Section
    doc.add_page_break()
    doc.add_heading('TEST SUMMARY', level=1)
    
    # Extract summary info from test_data
    doc.add_paragraph(f"Test ID: {test_id}")
    doc.add_paragraph(f"Total Steps: {len(frame_files)}")
    doc.add_paragraph(f"Status: ✓ Completed")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph()
    
    # Add AI analysis if available
    if test_data:
        doc.add_heading('Detailed Analysis', level=2)
        
        # Add only the relevant parts of AI response (skip frame-by-frame)
        lines = test_data.split('\n')
        skip_frames = False
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Skip the frame-by-frame section (we already have it above)
            if 'TEST STEPS' in line or '**Frame' in line:
                skip_frames = True
            elif line.startswith('##') and 'SUMMARY' in line:
                skip_frames = False
            
            if not skip_frames:
                if line.startswith('## '):
                    doc.add_heading(line.replace('## ', ''), level=2)
                elif line.startswith('# '):
                    continue  # Skip main title
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif not line.startswith('**Frame'):
                    doc.add_paragraph(line)
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    footer = doc.add_paragraph('Generated by Shakti-Gemini AI Assistant')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes

# Header
st.markdown('<p class="shakti-title">⚡ SHAKTI-GEMINI ⚡</p>', unsafe_allow_html=True)
st.markdown('<p class="subtitle">AI-Powered Assistant for Unit Testing & Data Analysis</p>', unsafe_allow_html=True)

# Sidebar
with st.sidebar:
    st.title("⚙️ Control Center")

    # Get API Key from Streamlit Secrets
    try:
        api_key = st.secrets["GEMINI_API_KEY"]
        st.success("✅ AI Ready!")
    except:
        st.error("⚠️ API key not configured. Contact admin.")
        api_key = None

    st.divider()

    # Choose Analysis Type
    st.subheader("🎯 Select Mode")
    
    analysis_mode = st.radio(
        "What do you want to analyze?",
        ["📹 Video (Unit Testing)", "📊 Excel File (Data Analysis)"],
        key="analysis_mode"
    )

    st.divider()

    # File Upload based on mode
    if analysis_mode == "📹 Video (Unit Testing)":
        st.subheader("🎥 Upload Test Video")
        
        if st.session_state.clear_uploader:
            st.session_state.clear_uploader = False
            st.session_state.uploader_key += 1

        uploaded_file = st.file_uploader(
            "Upload test execution video",
            type=['mp4', 'mov', 'avi', 'mkv', 'webm'],
            key=f"video_uploader_{st.session_state.uploader_key}",
        )

        if uploaded_file:
            st.video(uploaded_file)
            
            tfile = tempfile.NamedTemporaryFile(delete=False, suffix='.mp4')
            tfile.write(uploaded_file.read())
            video_path = tfile.name
            
            cap = cv2.VideoCapture(video_path)
            fps = cap.get(cv2.CAP_PROP_FPS)
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            duration = frame_count / fps if fps > 0 else 0
            cap.release()
            
            st.caption(f"⏱️ {duration:.1f}s | 🎞️ {fps:.0f} FPS")
            
            st.session_state.current_video_file = video_path
            st.session_state.video_duration = duration
            st.session_state.video_name = uploaded_file.name
        else:
            st.session_state.current_video_file = None

        st.divider()

        # Settings for video
        st.subheader("🎛️ Settings")
        
        num_screenshots = st.slider(
            "📸 Number of Screenshots:",
            min_value=4,
            max_value=12,
            value=8,
            help="Number of test steps to capture"
        )
        
        temperature = st.slider("🌡️ Detail Level:", 0.0, 1.0, 0.3)

    else:  # Excel Data mode
        st.subheader("📊 Upload Excel File")
        
        if st.session_state.clear_uploader:
            st.session_state.clear_uploader = False
            st.session_state.uploader_key += 1

        uploaded_file = st.file_uploader(
            "Upload Excel workbook (.xlsx or .xls)",
            type=['xlsx', 'xls'],
            key=f"excel_uploader_{st.session_state.uploader_key}",
            help="Upload Excel file with one or more sheets"
        )

        if uploaded_file:
            st.success(f"✅ File uploaded: {uploaded_file.name}")
            st.caption(f"📁 Size: {uploaded_file.size / 1024:.1f} KB")
            
            # Load all sheets
            with st.spinner("📖 Reading Excel sheets..."):
                sheets_dict, error = load_excel_sheets(uploaded_file)
            
            if error:
                st.error(f"❌ Error reading Excel file: {error}")
                st.session_state.excel_sheets = {}
            else:
                st.session_state.excel_sheets = sheets_dict
                
                # Show sheet info
                st.success(f"✅ Found {len(sheets_dict)} sheet(s)")
                for sheet_name, df in sheets_dict.items():
                    st.caption(f"📄 **{sheet_name}**: {len(df)} rows × {len(df.columns)} columns")
        else:
            st.session_state.excel_sheets = {}

        st.divider()

        # Analysis type for data
        st.subheader("🎛️ Analysis Settings")
        
        data_analysis_type = st.selectbox(
            "Select Analysis Type:",
            ["General Insights", "Stockout Risk", "Slow Movers"],
            help="Choose the type of analysis to perform"
        )
        
        # Map to prompt keys
        analysis_type_map = {
            "General Insights": "general_insights",
            "Stockout Risk": "stockout_risk",
            "Slow Movers": "slow_movers"
        }
        
        # Option to analyze all sheets or select specific ones
        if st.session_state.excel_sheets:
            st.subheader("📋 Select Sheets to Analyze")
            
            analyze_all = st.checkbox("Analyze all sheets", value=True)
            
            if not analyze_all and st.session_state.excel_sheets:
                selected_sheets = st.multiselect(
                    "Choose sheets:",
                    options=list(st.session_state.excel_sheets.keys()),
                    default=list(st.session_state.excel_sheets.keys())
                )
            else:
                selected_sheets = list(st.session_state.excel_sheets.keys())

    st.divider()

    if st.button("🗑️ Clear All", use_container_width=True):
        st.session_state.messages = []
        st.session_state.clear_uploader = True
        st.session_state.current_video_file = None
        st.session_state.unit_test_data = None
        st.session_state.captured_frames = []
        st.session_state.excel_sheets = {}
        st.session_state.sheet_analyses = {}
        st.rerun()

# Main Area
if uploaded_file or st.session_state.current_video_file:
    
    # Determine button text and action
    if analysis_mode == "📊 Excel File (Data Analysis)":
        button_text = "📊 ANALYZE EXCEL DATA"
        button_icon = "📊"
    else:
        button_text = "📝 GENERATE UNIT TEST DOCUMENT"
        button_icon = "📝"
    
    if st.button(button_text, type="primary", use_container_width=True):
        
        if not api_key:
            st.error("⚠️ API key not configured!")
            st.stop()
        
        try:
            # Excel Analysis Path
            if analysis_mode == "📊 Excel File (Data Analysis)" and st.session_state.excel_sheets:
                
                # Get sheets to analyze
                if 'selected_sheets' in locals():
                    sheets_to_analyze = {k: v for k, v in st.session_state.excel_sheets.items() if k in selected_sheets}
                else:
                    sheets_to_analyze = st.session_state.excel_sheets
                
                if not sheets_to_analyze:
                    st.warning("⚠️ No sheets selected for analysis!")
                    st.stop()
                
                total_sheets = len(sheets_to_analyze)
                
                # Progress tracking
                overall_progress = st.progress(0)
                status_text = st.empty()
                
                # Configure Gemini once
                genai.configure(api_key=api_key)
                model = genai.GenerativeModel("models/gemini-2.0-flash-exp")
                
                # Get analysis type
                selected_analysis = analysis_type_map[data_analysis_type]
                
                # Store analyses
                st.session_state.sheet_analyses = {}
                
                # Analyze each sheet
                for idx, (sheet_name, df) in enumerate(sheets_to_analyze.items(), 1):
                    
                    status_text.info(f"📊 Analyzing Sheet {idx}/{total_sheets}: **{sheet_name}**...")
                    overall_progress.progress(int((idx - 0.5) / total_sheets * 100))
                    
                    # Generate prompt for this sheet
                    full_prompt = analyze_sheet_data(df, sheet_name, selected_analysis)
                    
                    smart_rate_limit()
                    
                    # Get AI response
                    response = model.generate_content(
                        full_prompt,
                        generation_config={"temperature": 0.3},
                    )
                    
                    # Store the analysis
                    st.session_state.sheet_analyses[sheet_name] = {
                        'analysis': response.text,
                        'dataframe': df,
                        'rows': len(df),
                        'columns': len(df.columns)
                    }
                    
                    overall_progress.progress(int(idx / total_sheets * 100))
                
                overall_progress.empty()
                status_text.empty()
                
                # Display results
                st.success(f"✅ Analysis Complete! Analyzed {total_sheets} sheet(s)")
                
                # Summary metrics
                st.markdown("### 📊 Workbook Summary")
                cols = st.columns(4)
                
                total_rows = sum(info['rows'] for info in st.session_state.sheet_analyses.values())
                total_cols = sum(info['columns'] for info in st.session_state.sheet_analyses.values())
                
                cols[0].metric("Total Sheets", total_sheets)
                cols[1].metric("Total Rows", total_rows)
                cols[2].metric("Analysis Type", data_analysis_type)
                cols[3].metric("Generated", datetime.now().strftime("%H:%M"))
                
                st.divider()
                
                # Display each sheet's analysis
                for sheet_name, sheet_info in st.session_state.sheet_analyses.items():
                    
                    # Sheet header with custom styling
                    st.markdown(f'<div class="sheet-badge">📄 Sheet: {sheet_name}</div>', unsafe_allow_html=True)
                    
                    # Sheet metrics
                    col1, col2, col3 = st.columns(3)
                    col1.metric("Rows", sheet_info['rows'])
                    col2.metric("Columns", sheet_info['columns'])
                    col3.metric("Columns", ", ".join(sheet_info['dataframe'].columns.tolist()[:3]) + "...")
                    
                    # Data preview
                    with st.expander(f"📊 View Data - {sheet_name} (First 20 Rows)", expanded=False):
                        st.dataframe(sheet_info['dataframe'].head(20), use_container_width=True)
                    
                    # AI Analysis
                    st.markdown(f"### 🎯 AI Insights - {sheet_name}")
                    st.markdown(sheet_info['analysis'])
                    
                    # Download options for this sheet
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        # Download analysis
                        st.download_button(
                            label=f"📥 Download Analysis - {sheet_name}",
                            data=sheet_info['analysis'],
                            file_name=f"Analysis_{sheet_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                            mime="text/plain",
                            use_container_width=True,
                            key=f"download_analysis_{sheet_name}"
                        )
                    
                    with col2:
                        # Download data
                        csv_data = sheet_info['dataframe'].to_csv(index=False).encode('utf-8')
                        st.download_button(
                            label=f"📥 Download Data - {sheet_name}",
                            data=csv_data,
                            file_name=f"Data_{sheet_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                            mime="text/csv",
                            use_container_width=True,
                            key=f"download_data_{sheet_name}"
                        )
                    
                    st.divider()
                
                # Download combined report
                st.markdown("### 📥 Download Combined Report")
                
                combined_report = f"""
# EXCEL WORKBOOK ANALYSIS REPORT
Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
Analysis Type: {data_analysis_type}
Total Sheets Analyzed: {total_sheets}

---

"""
                for sheet_name, sheet_info in st.session_state.sheet_analyses.items():
                    combined_report += f"""
## 📄 SHEET: {sheet_name}

**Data Summary:**
- Rows: {sheet_info['rows']}
- Columns: {sheet_info['columns']}
- Column Names: {', '.join(sheet_info['dataframe'].columns.tolist())}

**AI Analysis:**

{sheet_info['analysis']}

{'=' * 80}

"""
                
                st.download_button(
                    label="📥 Download Complete Report (All Sheets)",
                    data=combined_report,
                    file_name=f"Complete_Analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain",
                    use_container_width=True
                )
                
                st.balloons()
            
            # Video Analysis Path (Original functionality)
            elif analysis_mode == "📹 Video (Unit Testing)" and st.session_state.current_video_file:
                
                # Progress tracking
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                # Step 1: Extract frames
                status_text.info(f"📸 Capturing {num_screenshots} screenshots...")
                progress_bar.progress(10)
                
                frames, frame_files = extract_video_frames_with_save(
                    st.session_state.current_video_file, 
                    num_screenshots
                )
                st.session_state.captured_frames = frame_files
                
                progress_bar.progress(30)
                
                # Step 2: AI Analysis
                status_text.info("🤖 Analyzing video with Gemini AI...")
                progress_bar.progress(40)
                
                genai.configure(api_key=api_key)
                model = genai.GenerativeModel("models/gemini-2.0-flash-exp")
                
                prompt = generate_unit_test_prompt()
                
                smart_rate_limit()
                
                # Use frames for faster analysis
                response = model.generate_content(
                    [prompt] + frames,
                    generation_config={"temperature": temperature},
                )
                
                progress_bar.progress(70)
                test_data = response.text
                st.session_state.unit_test_data = test_data
                
                # Step 3: Create document
                status_text.info("📝 Creating clean document...")
                progress_bar.progress(85)
                
                doc_bytes = create_clean_document_with_images(
                    test_data,
                    frame_files,
                    st.session_state.video_duration,
                    st.session_state.video_name
                )
                
                progress_bar.progress(100)
                progress_bar.empty()
                status_text.empty()
                
                st.success("✅ Clean Unit Test Document Generated!")
                
                # Preview
                with st.expander("👀 Preview Screenshots", expanded=False):
                    cols = st.columns(3)
                    for i, frame in enumerate(frames[:6]):
                        cols[i % 3].image(frame, caption=f"Step {i+1}", use_column_width=True)
                
                # Download
                st.download_button(
                    label="📥 DOWNLOAD UNIT TEST DOCUMENT (.docx)",
                    data=doc_bytes,
                    file_name=f"Unit_Test_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
                st.balloons()
            
            else:
                st.warning("⚠️ Please upload a file first!")
            
        except Exception as e:
            st.error(f"❌ Error: {str(e)}")
            st.exception(e)

else:
    # Welcome screen
    if analysis_mode == "📹 Video (Unit Testing)":
        st.info("👈 Upload a test execution video to begin")
        
        st.markdown("""
        ### ✨ Unit Test Generator Features:
        
        **📸 Large, Readable Screenshots**
        - 6-inch wide screenshots for clarity
        - Clear visibility of all UI elements
        
        **📝 Description Under Each Screenshot**
        - AI-generated description of what's happening
        - Centered, italicized text for clean look
        
        **🎯 Clean Document Layout**
        - Professional formatting
        - Easy to read and understand
        - Perfect for documentation and reports
        
        ### 📖 How to use:
        1. Upload your test execution video
        2. Choose number of steps (4-12)
        3. Click generate
        4. Download your clean document!
        """)
    
    else:  # Excel mode
        st.info("👈 Upload an Excel file to begin multi-sheet analysis")
        
        st.markdown("""
        ### ✨ Excel Multi-Sheet Analysis Features:
        
        **📊 Complete Workbook Analysis:**
        - Analyzes ALL sheets in your Excel file automatically
        - Each sheet gets its own dedicated AI analysis
        - Compare data across multiple sheets
        - Perfect for complex workbooks with multiple data sets
        
        **🎯 Three Analysis Types:**
        
        **1. General Insights**
        - Comprehensive overview of each sheet
        - Identifies data patterns and trends
        - Highlights key findings
        - Strategic recommendations
        
        **2. Stockout Risk Analysis**
        - Identifies products at risk per sheet
        - Calculates stockout dates
        - Recommends reorder quantities
        - Revenue impact assessment
        
        **3. Slow Movers Detection**
        - Finds excess inventory per sheet
        - Clearance recommendations
        - Cost savings analysis
        - Turnover optimization
        
        ### 📖 How to use:
        1. Upload your Excel workbook (.xlsx or .xls)
        2. App automatically detects all sheets
        3. Choose analysis type
        4. Select which sheets to analyze (or analyze all)
        5. Click analyze
        6. Get separate AI insights for EACH sheet!
        7. Download individual or combined reports
        
        ### 📋 Example Use Cases:
        - **Multi-location inventory**: Each sheet = different warehouse
        - **Product categories**: Each sheet = different category
        - **Time periods**: Each sheet = different month/quarter
        - **Departments**: Each sheet = different department data
        - **Suppliers**: Each sheet = different vendor data
        
        ### 💡 Pro Tips:
        - Name your sheets descriptively (e.g., "Warehouse_East", "Q4_2024")
        - Keep consistent column names across sheets for better analysis
        - Each sheet will be analyzed separately with full AI insights
        - You can compare insights across sheets easily
        """)
        
        # Sample Excel template info
        with st.expander("📥 Sample Excel Structure", expanded=False):
            st.markdown("""
            **Example Multi-Sheet Excel Structure:**
            
            **Sheet 1: Warehouse_North**
            | SKU | Product | Stock | Sales_30d | Category |
            |-----|---------|-------|-----------|----------|
            | SOF001 | Sofa A | 50 | 45 | Sofas |
            | TAB001 | Table B | 30 | 12 | Tables |
            
            **Sheet 2: Warehouse_South**
            | SKU | Product | Stock | Sales_30d | Category |
            |-----|---------|-------|-----------|----------|
            | SOF001 | Sofa A | 15 | 60 | Sofas |
            | BED001 | Bed C | 8 | 20 | Beds |
            
            **Sheet 3: Warehouse_West**
            | SKU | Product | Stock | Sales_30d | Category |
            |-----|---------|-------|-----------|----------|
            | CHA001 | Chair D | 100 | 5 | Chairs |
            | TAB002 | Table E | 25 | 18 | Tables |
            
            Each sheet will get its own detailed analysis!
            """)

st.divider()
st.caption(f"⚡ Shakti-Gemini AI Assistant | Multi-Sheet Excel Analysis | {datetime.now().strftime('%I:%M %p')}")
